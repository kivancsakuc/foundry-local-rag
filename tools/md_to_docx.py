"""Convert CURRICULUM.md (or any of this repo's markdown) to .docx for submission.

The repository keeps Markdown as the source of truth because it diffs; the internship
submission wants a Word document. This bridges the two so the docx is always generated,
never hand-edited.

    python tools/md_to_docx.py CURRICULUM.md
    python tools/md_to_docx.py CURRICULUM.md -o "Summer School Plan (revised).docx"

Handles the subset of Markdown this repo actually uses: ATX headings, paragraphs,
bullet and numbered lists, fenced code blocks, block quotes, pipe tables, horizontal
rules, and inline bold / italic / code / links.
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

# Inline markup, in the order it must be resolved. Links are unwrapped to their text
# followed by the URL, because a .docx reader cannot click a bare markdown link.
INLINE_RE = re.compile(
    r"(?P<code>`[^`]+`)"
    r"|(?P<bold>\*\*[^*]+\*\*)"
    r"|(?P<italic>(?<!\*)\*(?!\*)[^*]+\*(?!\*))"
    r"|(?P<link>\[[^\]]+\]\([^)]+\))"
    r"|(?P<autolink><https?://[^>]+>)"
)

CODE_COLOR = RGBColor(0xA0, 0x30, 0x60)


def add_inline(paragraph, text: str) -> None:
    """Write `text` into `paragraph`, converting inline markdown to real runs."""
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])

        kind = match.lastgroup
        raw = match.group()

        if kind == "code":
            run = paragraph.add_run(raw[1:-1])
            run.font.name = "Consolas"
            run.font.color.rgb = CODE_COLOR
        elif kind == "bold":
            paragraph.add_run(raw[2:-2]).bold = True
        elif kind == "italic":
            paragraph.add_run(raw[1:-1]).italic = True
        elif kind == "link":
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", raw).groups()
            paragraph.add_run(label)
            # Keep the destination visible; a generated docx has no live hyperlinks.
            if not label.startswith("http"):
                small = paragraph.add_run(f" ({url})")
                small.font.size = Pt(8)
                small.italic = True
        elif kind == "autolink":
            paragraph.add_run(raw[1:-1])

        pos = match.end()

    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_code_block(doc: Document, lines: list[str]) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Pt(18)
    para.paragraph_format.space_after = Pt(10)
    run = para.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    """Render a pipe table. The first row is the header."""
    if not rows:
        return
    width = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=width)
    table.style = "Light Grid Accent 1"

    for r, row in enumerate(rows):
        for c in range(width):
            cell = table.cell(r, c)
            cell.text = ""
            para = cell.paragraphs[0]
            add_inline(para, row[c] if c < len(row) else "")
            if r == 0:
                for run in para.runs:
                    run.bold = True
    doc.add_paragraph()


def split_table_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_table_divider(line: str) -> bool:
    return bool(re.fullmatch(r"\s*\|?[\s:|-]+\|?\s*", line)) and "-" in line


def starts_block(line: str) -> bool:
    """True if `line` begins a new markdown block rather than continuing one."""
    return (
        line.startswith(("#", ">", "|", "```"))
        or bool(re.match(r"[-*+]\s+", line))
        or bool(re.match(r"\d+\.\s+", line))
        or bool(re.fullmatch(r"-{3,}|\*{3,}|_{3,}", line))
    )


def convert(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Fenced code block
        if stripped.startswith("```"):
            i += 1
            block: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1
            add_code_block(doc, block)
            continue

        # Pipe table: a header row followed by a divider row
        if (
            stripped.startswith("|")
            and i + 1 < len(lines)
            and is_table_divider(lines[i + 1])
        ):
            rows = [split_table_row(stripped)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_table_row(lines[i]))
                i += 1
            add_table(doc, rows)
            continue

        # Horizontal rule
        if re.fullmatch(r"-{3,}|\*{3,}|_{3,}", stripped):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run("* * *")
            run.font.size = Pt(9)
            i += 1
            continue

        # Heading
        heading = re.match(r"(#{1,6})\s+(.*)", stripped)
        if heading:
            level = len(heading.group(1))
            para = doc.add_heading(level=min(level, 4))
            para.text = ""
            add_inline(para, heading.group(2))
            i += 1
            continue

        # Block quote
        if stripped.startswith(">"):
            quote: list[str] = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote.append(lines[i].strip().lstrip(">").strip())
                i += 1
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Pt(18)
            add_inline(para, " ".join(q for q in quote if q))
            for run in para.runs:
                run.italic = True
            continue

        # List item, bullet or numbered. A wrapped item continues on the following
        # lines until a blank line or the start of another block, so those lines must
        # be folded in rather than becoming stray paragraphs.
        bullet = re.match(r"[-*+]\s+(.*)", stripped)
        numbered = re.match(r"\d+\.\s+(.*)", stripped)
        if bullet or numbered:
            match = bullet or numbered
            style = "List Bullet" if bullet else "List Number"
            item = [match.group(1)]
            i += 1
            while i < len(lines):
                nxt = lines[i].strip()
                if not nxt or starts_block(nxt):
                    break
                item.append(nxt)
                i += 1
            para = doc.add_paragraph(style=style)
            add_inline(para, " ".join(item))
            continue

        # Blank line
        if not stripped:
            i += 1
            continue

        # Paragraph: join continuation lines so wrapped Markdown does not become
        # one Word paragraph per source line.
        body = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or starts_block(nxt):
                break
            body.append(nxt)
            i += 1

        para = doc.add_paragraph()
        add_inline(para, " ".join(body))

    doc.save(docx_path)


def main(argv: list[str]) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("markdown", type=Path, help="source .md file")
    parser.add_argument("-o", "--output", type=Path, help="destination .docx file")
    args = parser.parse_args(argv[1:])

    if not args.markdown.exists():
        print(f"Not found: {args.markdown}", file=sys.stderr)
        return 1

    output = args.output or args.markdown.with_suffix(".docx")
    convert(args.markdown, output)
    print(f"Wrote {output} ({output.stat().st_size:,} bytes)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
